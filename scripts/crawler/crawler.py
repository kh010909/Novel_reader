# Web crawler :小說狂人(fiction maniac)
import requests  # requests tutorial: https://steam.oxxostudio.tw/category/python/spider/requests.html
import chardet #自動檢測網頁編碼:check jjwxc coding
import pymysql  # py connect mysql: https://www.learncodewithmike.com/2020/02/python-mysql.html
import os
import sys
from opencc import OpenCC # chinese conversion: https://yanwei-liu.medium.com/python%E8%87%AA%E7%84%B6%E8%AA%9E%E8%A8%80%E8%99%95%E7%90%86-%E5%9B%9B-%E7%B9%81%E7%B0%A1%E8%BD%89%E6%8F%9B%E5%88%A9%E5%99%A8opencc-74021cbc6de3
import bs4  # bs4 tutorial: https://www.learncodewithmike.com/2020/02/python-beautifulsoup-web-scraper.html、
#                           https://seanchien0525.medium.com/python-requests-beautifulsoup-%E7%88%AC%E8%9F%B2%E6%95%99%E5%AD%B8-83d146faa9e8
# basic crawler tutorial: https://hackmd.io/@AndyChiang/StaticCrawler
from docx import Document
import time
# from flask import Flask, request, send_file, redirect, url_for
# from io import BytesIO
#動態網頁爬蟲
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.chrome.options import Options
import bgm

# connect mysql
conn = pymysql.connect(host='localhost', port=3306, user='root',
                       passwd='', charset='utf8', db='novel_reader')

#chinese conversion
cc_s2tw = OpenCC('s2tw') #Simplified Chinese to Traditional Chinese

def get_czbooks_novel(url):
    # Create a request object, then use the object to open the URL, and enter the request headers information
    data = requests.get(url)
    if data.status_code!=200:
        print("Fail to connect to the website you entered.")
        print("Please make sure the url you input.")
        exit(1)
    # Parse the source code and get the title of each article
    root = bs4.BeautifulSoup(data.text, "html.parser")
    nName = root.find("span", class_="title")
    if nName==None:
        print("The URL you input is error")
        exit(1)
    nName = nName.string
    # Remove characters at the beginning or end of a string:https://www.freecodecamp.org/chinese/news/python-strip-how-to-trim-a-string-or-line/
    nName = nName.strip("《》")
    # print("novel name:"+nName)
    author = root.find("span", class_="author")
    author = author.select_one("a")
    author = author.string
    # print(author)

    description = root.find("div", class_="description")
    description = description.decode_contents()
    # print(description)

    # If select or find cannot find the corresponding content, it will respond:[]
    if (root.select(".thumbnail-state", limit=1) != []):
        completed = "完結"
    else:
        completed = "連載"

    # tags
    tags_lable = root.find("ul", class_="hashtag")
    tags_lable = tags_lable.select(".item")
    tags=[cc_s2tw.convert(tag.a.string) for tag in tags_lable]
    
    # img url
    nImg_link = root.find("div", class_="thumbnail")
    nImg_link = nImg_link.find("img")
    nImg_link = nImg_link['src']

    # 抓取下一頁的資料
    nextLink = root.find("li", class_="volume")
    if nextLink is None:
        print("The nextLink can't be found")
        exit(1)
    nextLink = nextLink.find_next_sibling("li")
    nextLink = nextLink.a
    nextLink = "https:"+nextLink["href"]
    return nName, author, description, completed, tags, nImg_link, nextLink


def get_czbooks_article(url):
    # Create a request object, then use the object to open the URL, and enter the request headers information
    data = requests.get(url)
    if data.status_code!=200:
        print("Fail to connect to web")
        exit(1)
    # Parse the source code and get the title of each article
    root = bs4.BeautifulSoup(data.text, "html.parser")
    aName = root.find("div", class_="name")
    aName = aName.string
    
    aContent = root.find("div", class_="content")
    aContent = aContent.decode_contents()

    # Fetch the next page of data
    nextLink = root.find("a", class_="next-chapter")
    if nextLink is not None:
        nextLink = "https:"+nextLink["href"]
    # print(nextLink)

    return aName, aContent, nextLink


def get_qidian_novel(url):
    # Create a request object, then use the object to open the URL, and enter the request headers information
    data = requests.get(url)
    if data.status_code!=200:
        print("Fail to connect to the website you entered.")
        print("Please make sure the url you input.")
        exit(1)
    # Parse the source code and get the title of each article
    root = bs4.BeautifulSoup(data.text, "html.parser")
    nName = root.find("meta", property="og:novel:book_name")
    if nName==None:
        print("The URL you input is error")
        exit(1)
    nName = nName["content"]
    # print("novel name:"+nName)
    
    author = root.find("meta", property="og:novel:author")
    author = author["content"]

    description = root.find("div", class_="book-intro")
    description = description.select_one("p")
    des=description.decode_contents()
    # print(des)

    # # If select or find cannot find the corresponding content, it will respond:[]
    completed = root.find("meta", property="og:novel:status")
    completed = completed["content"]

    # # tags
    tags_lable = root.find("p", class_="tag")
    tags_lable = tags_lable.select("span")+tags_lable.select("a")

    if (completed == "完本"):  # "完本"轉成"完結"
        completed = "完結"
        tags_lable[0].string = "完結"
    

    # print("completed:", completed)
    tags=[cc_s2tw.convert(tag.string) for tag in tags_lable]
    

    # img url
    nImg_link = root.find("meta", property="og:image")
    nImg_link = "https:"+nImg_link["content"]
    # print(nImg_link)

    # 抓取下一頁的資料
    nextLink = root.find("a", id="readBtn")
    if nextLink is None:
        print("The article link can't be found")
        exit(1)
    nextLink = "https:"+nextLink["href"]
    # print(nextLink)
    
    return cc_s2tw.convert(nName), cc_s2tw.convert(author), cc_s2tw.convert(des), cc_s2tw.convert(completed), tags, nImg_link, nextLink


def get_qidian_article(url):
    
    # Create a request object, then use the object to open the URL, and enter the request headers information
    data = requests.get(url)
    if data.status_code!=200:
        print("Fail to connect to web")
        exit(1)
    time.sleep(1)
    # Parse the source code and get the title of each article
    root = bs4.BeautifulSoup(data.text, "html.parser")
    aName = root.find("span", class_="content-wrap")
    if aName != None:
        aName = aName.string
        # print("article name(1):"+str(aName))

        aContent = root.find("div", class_="read-content j_readContent")
        aContent = aContent.decode_contents()
        # print(aContent)

        # Fetch the next page of data
        nextLink = root.find("a", string="下一章")
        if nextLink is not None:
            nextLink = "https:"+nextLink["href"]
        if nextLink is not None and "vipreader" in nextLink:
            nextLink=None
        # print("nextlink(1):",nextLink) 
    else:
        aName=root.find("meta", attrs={"name": "keywords"})
        aName=aName["content"].split(",")
        aName=aName[1]
        # print("article name(2):"+str(aName))
        aContent=root.find("main", attrs={"data-type": "cjk"})
        aContent=aContent.decode_contents()
        # print(aContent)
        nextLink = root.find_all("a", attrs={"class": "nav-btn"})
        if len(nextLink)==3:
            nextLink=nextLink[2]
        elif len(nextLink)==2:
            nextLink=nextLink[1]
        if nextLink is not None:
            nextLink = "https:"+nextLink["href"]
        if nextLink is not None and ("vipreader" in nextLink or "lastpage" in nextLink):
            nextLink=None
        # print("nextlink(2):",nextLink)

    return cc_s2tw.convert(aName), cc_s2tw.convert(aContent), nextLink

def find_index_of_keyword(lst):
    keyword = "主角"
    for i, item in enumerate(lst):
        if keyword in item:
            return i
        
    return -1  # 如果列表中没有出现关键字，则返回 -1

def get_jjwxc_novel(url):
    # Parse webpage and convert text encoding from gb18030 to utf-8
    data = requests.get(url) #get wb content
    if data.status_code!=200:
        print("Fail to connect to the website you entered.")
        print("Please make sure the url you input.")
        exit(1)
    data_transcoded = data.content.decode('gb18030', errors='replace').encode('utf-8').decode('utf-8')
    root=bs4.BeautifulSoup(data_transcoded, "html.parser")
    
    nName = root.find("span", itemprop="articleSection")
    if nName==None:
        print("The URL you input is error")
        exit(1)
    nName = nName.string
    # print("novel name:"+nName)
    
    # author = root.find("span", itemprop="author")
    # author = author.string
    keywords = root.find("meta", attrs={'name':'Description'})
    keywords=keywords["content"].split(',')
    author=keywords[1]
    # print("author:"+author)

    description = root.find("div", id="novelintro")
    des = description.decode_contents()
    # print(des)
    # print("description:"+des)

    completed = root.find("span", itemprop="updataStatus")
    completed = completed.string
    # print("completed:"+completed)

    # tags
    ind=find_index_of_keyword(keywords)
    if ind==-1:
        tags_tag=keywords[2].split()
    else:
        tags_tag=keywords[ind-1].split()
        
        tags_charater=keywords[ind].split(' ┃ ')
        
        if "主角：" in tags_charater[0]:
            temp=tags_charater[0].replace("主角：", "")
            # print(temp)
            tags_main_charaters=temp.split('，')
            tags_simplified=tags_tag+tags_main_charaters
        else:
            tags_simplified=tags_tag
    
    tags=[cc_s2tw.convert(tag) for tag in tags_simplified]
    # for tag in tags:
    #     print("novel tag:"+tag)

    # img url
    nImg_link = root.find("img", class_="noveldefaultimage")
    nImg_link = nImg_link["src"]
    # print(nImg_link)

    # 抓取下一頁的資料
    nextLink = root.find("a", itemprop="url")
    if nextLink is None:
        print("The article link can't be found")
        exit(1)
    nextLink = nextLink["href"]
    # print(nextLink)

    return cc_s2tw.convert(nName), cc_s2tw.convert(author), cc_s2tw.convert(des), cc_s2tw.convert(completed), tags, nImg_link, nextLink


def get_jjwxc_article(url, driver):
    
    # 打开网页
    driver.get(url)
    # 设置显式等待时间为15秒
    wait = WebDriverWait(driver, 15)
    
    try:
        # 使用了显式等待（Explicit Wait）来等待页面加载完成，使用 wait 对象调用 until 方法，并传递一个条件（Expected Condition）作为参数，
        #   EC.presence_of_element_located((By.TAG_NAME, 'body'))，它表示等待页面中的 <body> 标签元素出现
        # 等待页面加载完成
        page_loaded = wait.until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
    except TimeoutException:
        print("页面加载超时")
        exit(1)
        # return None, None, None, aChapter
    
    # 获取页面内容
    root=bs4.BeautifulSoup(driver.page_source, "html.parser")
    
    # 进一步处理页面内容...
    # print("爬蟲的網址:", driver.current_url)
    aName_with_tag = root.find("div", style="clear:both;")
    if aName_with_tag is None:
        lockpage=root.find("div", id="lockpage")
        if lockpage is None:
            print("The novel has VIP article")
            return None, None, None
        else:
            lockpage=lockpage.select("p")
            lockpage=lockpage[8].find_all("a", style="cursor:pointer")
            nextLink="https://www.jjwxc.net"+lockpage[1]["href"]
            # print("被鎖的網址:", nextLink)
            return "[鎖]", "此章節為原網站的被鎖章節", nextLink
    aName_with_tag = aName_with_tag.find_previous_sibling("div")
    aName = aName_with_tag.select_one("h2")
    aName = aName.string
    # print("article name:"+aName)

    aContent_target = root.find("div", style="clear:both;")
    aContent_next_target = aContent_target.find_next_sibling(string=True)
    next_br=aContent_next_target.find_next_sibling("br")
    aContent=""
    while aContent_next_target and next_br and aContent_next_target.name!="div":
        aContent+=str(aContent_next_target.strip())
        aContent+=str(next_br)
        next_br=aContent_next_target.find_next_sibling("br")
        aContent_next_target=aContent_next_target.find_next_sibling(string=True)

    # Fetch the next page of data
    nextLink = root.find("a", string="下一章→")
    if nextLink is None:
        print("The nextLink can't be found")
        exit(1)
    if "noveloverlist.php" in nextLink["href"]: #the last page
        nextLink=None
    else:
        nextLink="https://www.jjwxc.net/"+nextLink["href"]
    # if "chapterid" not in nextLink:
    #     nextLink=None
    # print(nextLink) 
        
    return cc_s2tw.convert(aName), cc_s2tw.convert(aContent), nextLink
    

# Insert novel image
def imgToDB(nImg_link, nId, in_or_up):  # in_or_up:the novel is inserted(1) or updated(0)
    # local storage path(initial_crawler.py)
    dir = "../../static/images/novel/"
    # py使用檔案系統:https://medium.com/seaniap/python-%E4%BD%BF%E7%94%A8%E6%AA%94%E6%A1%88%E7%B3%BB%E7%B5%B1-eaecbe7001dd, https://ithelp.ithome.com.tw/m/articles/10262264
    # vscode中使用python相对路径问题: https://www.zhihu.com/question/313379182
    # modify current work list and find the current location of initial_crawler.py
    os.chdir(sys.path[0])
    # print(os.getcwd())  # print modified work list
    # make sure whether the default img
    if (nImg_link != "https://czbooks.net/images/default_no_thumbnail.jpg"):
        # Check if the local path exists
        if not os.path.exists(dir):
            os.mkdir(dir)  # build file
        # Use a crawler to download images:https://www.learncodewithmike.com/2020/09/download-images-using-python.html
        response = requests.get(nImg_link)  # download img
        if response.status_code == 200:  # write into file
            nId = "".join(str(id) for id in nId)  # convert tuple into string
            nImg = "novel_"+nId+".jpg"
            download_dir = dir+nImg
            # if novel is updated and original img exists
            if (in_or_up == 0 and os.path.exists(download_dir)):
                # remove original file:https://www.cjkuo.net/python_delete_file_or_dir/
                os.remove(download_dir)
            with open(download_dir, 'wb') as f:
                f.write(response.content)
            # print('Image download successful')
        else:
            print("Image download failed, error code: ", response.status_code)
            exit(1)
        f.close()    # close file
        with conn.cursor() as cursor:
            sql = "UPDATE novel SET nImg = %s WHERE nId = %s"
            cursor.execute(sql, (nImg, nId))  # execute sql command
            conn.commit()  # submit to sql
    else:
        if (in_or_up == 0):  # update novel
            with conn.cursor() as cursor:
                sql = "UPDATE novel SET nImg = %s WHERE nId = %s"
                cursor.execute(sql, ("default.jpg", nId))
                conn.commit()

                # remove original img
                nId = "".join(str(id)
                              for id in nId)  # convert tuple into string
                nImg = "novel_"+nId+".jpg"
                download_dir = dir+nImg
                if os.path.isfile(download_dir):
                    os.remove(download_dir)  # remove original file
        # print("默認圖片")


# novel table: nId, nName, author, nImg, description, completed(Is it over), nTime, nLike(Likes)
# tag  table: nId, tag
def novelToDB(nName, author, description, completed, tags, nImg_link, nextLink):

    # Insert into DB
    with conn.cursor() as cursor:
        # Insert if it does not exist, do not insert if it exists: https://blog.csdn.net/u012373815/article/details/75062444
        sql = "INSERT INTO novel(nName, author, description, completed) \
        SELECT %s, %s, %s, %s \
        FROM dual \
        WHERE NOT EXISTS(SELECT nName FROM novel WHERE nName=%s AND author=%s);"
        # dual is a virtual table
        cursor.execute(sql, (nName, author, description,
                       completed, nName, author))
        # print("novel插入還是不插入:", cursor.rowcount)
        if cursor.rowcount == 0:  # not insert
            sql = "UPDATE novel SET description=%s, completed=%s WHERE  nName=%s AND author=%s;"
            cursor.execute(sql, (description, completed, nName, author))
            # print("novel有沒有update:", cursor.rowcount)
            conn.commit()

            # get nId
            sql = "SELECT nId FROM novel WHERE nName=%s AND author=%s"
            cursor.execute(sql, (nName, author))
            nId = cursor.fetchone()

            # update image
            imgToDB(nImg_link, nId, 0)

        elif cursor.rowcount == 1:  # insert
            conn.commit()

            # get nId
            sql = "SELECT nId FROM novel WHERE nName=%s AND author=%s"
            cursor.execute(sql, (nName, author))
            nId = cursor.fetchone()
            # print(nId)
            # insert tag
            if tags != []:
                for tag in tags:
                    # print(tag)
                    if tag not in ["已完結", "未完結", "完結", "斷章", "連載"] :
                        sql = "INSERT INTO tag(nId, tag) \
                            SELECT %s, %s \
                            FROM dual \
                            WHERE NOT EXISTS(SELECT tag FROM tag WHERE nId=%s AND tag=%s);"
                        cursor.execute(sql, (nId, tag, nId, tag))
                        # if(cursor.rowcount==0):
                        #     print("tag not insert")
                        # else:
                        #     print("tag insert")
                        conn.commit()
                        

            # main tag must be inserted to DB
            # sql = "INSERT INTO tag(nId) VALUES(%s)"
            # cursor.execute(sql, (nId))
            # conn.commit()

            # insert image
            imgToDB(nImg_link, nId, 1)

    # 回傳下一頁資料
    # print(nextLink)
    return nextLink, nId

# Article: nId, aId, aName, aChapter, aContent, mId
def articleToDB(aChapter, nId, aName, aContent, nextLink, addBgm):
    if addBgm=="yes":
        aContent= bgm.add_music(aContent)
        # print(aContent)

    with conn.cursor() as cursor:
        sql = "INSERT INTO article(nId, aName, aChapter, aContent, mId) \
        SELECT %s, %s, %s, %s, %s \
        FROM dual \
        WHERE NOT EXISTS(SELECT nId, aChapter FROM article WHERE nId=%s AND aChapter=%s);"
        # dual is virtual table
        cursor.execute(sql, (nId, aName, aChapter, aContent,
                       1, nId, aChapter))
        # print("article插入還是不插入:", cursor.rowcount)
        if cursor.rowcount == 0:
            sql = "UPDATE article SET aName=%s, aContent=%s WHERE  nId=%s AND aChapter=%s;"
            cursor.execute(sql, (aName, aContent, nId, aChapter))
            # print("article有沒有update:", cursor.rowcount)

        conn.commit()

    if (nextLink == None):  # The last page of the article, update novel's newestChapter and nTime
        with conn.cursor() as cursor:
            sql = "SELECT aTime FROM article WHERE aName=%s AND aChapter=%s"
            cursor.execute(sql, (aName, aChapter))
            aTime = cursor.fetchone()
            sql = "UPDATE novel SET newestChapter=%s, nTime=%s WHERE nId=%s"
            cursor.execute(sql, (aName, aTime, nId))
            conn.commit()

        return nextLink
    else:
        return nextLink

def crawlerToDB(pageURL, wb_name, addBgm):
    
    # insert novel
    #wb_name include: 小說狂人(czbooks)、起點中文網(qidian)、晉江文學城(jjwxc)
    if (wb_name == "czbooks"):
        pageURL= "https://czbooks.net/n/"+pageURL
        nName, author, description, completed, tags, nImg_link, nextLink = get_czbooks_novel(
            pageURL)
    elif (wb_name == "qidian"):
        pageURL= "https://book.qidian.com/info/"+pageURL
        nName, author, description, completed, tags, nImg_link, nextLink = get_qidian_novel(
            pageURL)
    elif (wb_name == "jjwxc"):
        pageURL= "https://www.jjwxc.net/onebook.php?novelid="+pageURL
        nName, author, description, completed, tags, nImg_link, nextLink = get_jjwxc_novel(
            pageURL)
    articleURL, nId = novelToDB(nName, author, description, completed,
                                tags, nImg_link, nextLink)  # (page url)
    pageURL = articleURL

    # insert article
    
    nextURL = "0"
    if(wb_name == "czbooks"):
        aChapter = 1
        while nextURL != "":
            aName, aContent, nextLink = get_czbooks_article(pageURL)

            nextURL = articleToDB(aChapter, nId, aName, aContent, nextLink, addBgm)
            # 防止TypeError: can only concatenate str (not "NoneType") to str
            if nextURL is None:
                nextURL = ""
            pageURL = nextURL
            aChapter += 1
    elif(wb_name=="qidian"):
        aChapter = 1
        while nextURL != "":
            aName, aContent, nextLink = get_qidian_article(pageURL)
            aContent=aContent.replace("<p>", "<br/>")
            aContent=aContent.replace("</p>", "")
            nextURL = articleToDB(aChapter, nId, aName, aContent, nextLink, addBgm)
            # 防止TypeError: can only concatenate str (not "NoneType") to str
            if nextURL is None:
                nextURL = ""
            pageURL = nextURL
            aChapter += 1
    elif(wb_name=="jjwxc"): 
        aChapter = 1
        options = Options()
        options.add_argument("--disable-notifications") #用于禁用浏览器的通知功能
        options.add_argument("--headless")  # 启用无头模式，无头模式是一种在后台运行浏览器的方式，没有可见的图形用户界面
        options.add_argument("--disable-gpu")  # 禁用 GPU 加速，避免在某些系统上可能出现的 GPU 相关问题
        options.add_experimental_option('excludeSwitches', ['enable-logging']) #關閉DevTools listening 的打印(包含警告)
        # 创建浏览器对象
        driver = webdriver.Chrome('./chromedriver', options=options)
        while nextURL != "":
            aName, aContent, nextLink = get_jjwxc_article(pageURL, driver)
            if not aName and not aContent and not nextLink :
                aChapter-=1
                with conn.cursor() as cursor:
                    sql = "SELECT aTime, aName FROM article WHERE nId=%s AND aChapter=%s"
                    cursor.execute(sql, (nId, aChapter))
                    aTime, aName = cursor.fetchone()
                    sql = "UPDATE novel SET newestChapter=%s, nTime=%s WHERE nId=%s"
                    cursor.execute(sql, (aName, aTime, nId))
                    conn.commit()
                break
            
            nextURL = articleToDB(aChapter, nId, aName, aContent, nextLink, addBgm)
            # 防止TypeError: can only concatenate str (not "NoneType") to str
            if nextURL is None:
                nextURL = ""
            pageURL = nextURL
            aChapter += 1
        driver.quit()
    conn.close()
    return "The novel is available to read on this website !"

def crawlerToDocx(pageURL, wb_name):
    # insert novel
    #wb_name include: 小說狂人(czbooks)、起點中文網(qidian)、晉江文學城(jjwxc)
    if (wb_name == "czbooks"):
        pageURL= "https://czbooks.net/n/"+pageURL
        nName, author, description, completed, tags, nImg_link, nextLink = get_czbooks_novel(
            pageURL)
        nName=nName.replace("<br>", "")
        author=author.replace("<br>", "")
        description=description.replace("<br/>","\n")
        description=description.replace("</br>","\n")
        description=description.replace("<br>","\n")
    elif (wb_name == "qidian"):
        pageURL= "https://book.qidian.com/info/"+pageURL
        nName, author, description, completed, tags, nImg_link, nextLink = get_qidian_novel(
            pageURL)
        description=description.replace("<br>","\n")
        description=description.replace("<br/>","\n")
    elif (wb_name == "jjwxc"):
        pageURL= "https://www.jjwxc.net/onebook.php?novelid="+pageURL
        nName, author, description, completed, tags, nImg_link, nextLink = get_jjwxc_novel(
            pageURL)
        description=description.replace("<br>","\n")
        description=description.replace("<br/>","\n")
    
    pageURL = nextLink
    # 創建一個新的Word文檔
    document = Document()

    # 添加標題
    document.add_heading(nName, level=0)
    # 添加段落
    document.add_paragraph("作者:"+author)
    document.add_paragraph("狀態:"+completed)
    document.add_paragraph("作品簡介:\n\t"+description)
    document.add_paragraph("標籤:"+",".join(tags))

    # insert article
    if(wb_name == "czbooks"):
        aChapter = 0
        while nextLink != "":
            aName, aContent, nextLink = get_czbooks_article(pageURL)
            aName=aName.replace("<br>","")
            aContent=aContent.replace("<br/>", "\u2028")
            aContent=aContent.replace("</br>", "\u2028")
            aContent=aContent.replace("<br>", "\u2028")
            document.add_paragraph("第"+str(aChapter)+"章\n"+aName+"\n"+aContent+"\n")
            
            # 防止TypeError: can only concatenate str (not "NoneType") to str
            if nextLink is None:
                nextLink = ""
            pageURL = nextLink
            aChapter += 1  
    elif(wb_name=="qidian"):
        aChapter = 1
        while nextLink != "":
            aName, aContent, nextLink = get_qidian_article(pageURL)
            aContent=aContent.replace("<p>", "\n")
            aContent=aContent.replace("</p>", "")
            document.add_paragraph("第"+str(aChapter)+"章\u2028"+aName+"\n"+aContent+"\n")
            
            # 防止TypeError: can only concatenate str (not "NoneType") to str
            if nextLink is None:
                nextLink = ""
            pageURL = nextLink
            aChapter += 1
    elif(wb_name=="jjwxc"): 
        aChapter = 1
        options = Options()
        options.add_argument("--disable-notifications") #用于禁用浏览器的通知功能
        options.add_argument("--headless")  # 启用无头模式，无头模式是一种在后台运行浏览器的方式，没有可见的图形用户界面
        options.add_argument("--disable-gpu")  # 禁用 GPU 加速，避免在某些系统上可能出现的 GPU 相关问题
        options.add_experimental_option('excludeSwitches', ['enable-logging']) #關閉DevTools listening 的打印(包含警告)
        # 创建浏览器对象
        driver = webdriver.Chrome('./chromedriver', options=options)
        while nextLink != "":
            aName, aContent, nextLink = get_jjwxc_article(pageURL, driver)
            if not aName and not aContent and not nextLink :
                document.add_paragraph("進入VIP章節")
            else:
                aContent=aContent.replace("<br>", "\n")
                aContent=aContent.replace("<br/>", "\n")
                document.add_paragraph("第"+str(aChapter)+"章\u2028"+aName+"\n"+aContent+"\n")
            
            # 防止TypeError: can only concatenate str (not "NoneType") to str
            if nextLink is None:
                nextLink = ""
            pageURL = nextLink
            aChapter += 1
        driver.quit()

    #Store to local
    dir="../../static/novel"
    file_name=repr(nName).replace("/", "_")
    file_name=file_name.strip("'")+".docx"
    file_path = os.path.abspath(os.path.join(dir, file_name))
    os.chdir(sys.path[0])
    if not os.path.exists(dir):
        os.mkdir(dir)  # build file
        # print("build file.")
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            # print("Remove existing file")
        except PermissionError:
            exit(1)
        
    # 儲存文檔
    document.save(file_path)
    return file_path

# main funciton
# wb_name:小說狂人(czbooks)、起點中文網(qidian)、晉江文學城(jjwxc)
# 接收命令行参数
functionName = sys.argv[1]
pageURL = sys.argv[2]
wb_name = sys.argv[3]
if len(sys.argv)>=5:
    addBgm = sys.argv[4]
    if functionName=="crawlerToDB":
        result=crawlerToDB(pageURL, wb_name, addBgm)
else:
    if functionName=="crawlerToDocx":
        result=crawlerToDocx(pageURL, wb_name)
# 返回结果
print(result)

# get_qidian_article("https://www.qidian.com/chapter/1034130287/751765931/")
#測試:
# result=crawlerToDocx("3694555", "jjwxc")
# result=crawlerToDB("1037017211/", wb_name="qidian", addBgm="yes")
# print(result)
# pageURL = "3694555"

#/xampp/htdocs/Novel_reader/scripts/crawler/crawler.py
# wb_name = request.form['wb_name']
# pageURL = request.form['czbooks-url']
# print(wb_name)
#wb_name:小說狂人(czbooks)、起點中文網(qidian)、晉江文學城(jjwxc)
# 初始化網址
# pageURL = "https://czbooks.net/n/cpgb0a1"
# pageURL = "https://book.qidian.com/info/1031931986/#Catalog" #末日沒來
# pageURL = "https://book.qidian.com/info/1017420102/#Catalog" #收账(has vip章節)
# pageURL = "https://www.jjwxc.net/onebook.php?novelid=2433649" #護心
# pageURL = "https://www.jjwxc.net/onebook.php?novelid=3519615" #哭泣的天使
# crawlerToDB(pageURL, wb_name="jjwxc")
# get_czbooks_novel(pageURL)
# get_czbooks_article(pageURL)

# crawler to local
# crawlerToDocx(pageURL, "jjwxc", "C:/xampp/htdocs/Novel_reader/scripts/crawler/") #pageURL, wb_name, dir
# crawlerToHtml(pageURL, "jjwxc", "C:/xampp/htdocs/Novel_reader/scripts/crawler/") #pageURL, wb_name, dir
# localToWeb("毒液/Venom", "C:/xampp/htdocs/Novel_reader/scripts/crawler/") #nName, dir

# pageURL = "https://book.qidian.com/info/1031931986/#Catalog" #末日沒來
# pageURL = "https://book.qidian.com/info/1017420102/#Catalog" #收账(has vip章節)
# crawlerToDB(pageURL, main_tag="短篇", wb_name="qidian")
# get_qidian_novel(pageURL)
# pageURL = "https://vipreader.qidian.com/chapter/1017420102/511368442/"
# get_qidian_article(pageURL)

# pageURL = "https://www.jjwxc.net/onebook.php?novelid=3519615" #哭泣的天使
# pageURL ="https://www.jjwxc.net/onebook.php?novelid=6229060" #男友死在三年前(has vip章節)
# pageURL = "https://www.jjwxc.net/onebook.php?novelid=2433649" #護心
# crawlerToDB(pageURL, main_tag="短篇", wb_name="jjwxc")
# get_jjwxc_novel(pageURL)
# pageURL ="https://www.jjwxc.net/onebook.php?novelid=6229060&chapterid=23"
# pageURL ="https://www.jjwxc.net/onebook.php?novelid=6711365&chapterid=34"
# pageURL = "https://www.jjwxc.net/onebook.php?novelid=4174255&chapterid=1"
# get_jjwxc_article(pageURL)


# 已解決和待解決問題
# article的最新時間更新到novel的nTime(ok)
# 避免同樣的小說被重複insert進DB(ok)
# 小說的連載、完結、斷更(找不到斷更的判斷依據)(ok)
# 等music的DB完成，要把article的mId(foreign key)和music的mId連結----------
# 修改article update的部分(用aChapter, nId當判斷基準，決定是update還是insert)(ok)
# 檢查novel裡，如果data被修改了，是否會update(ok)
# img也可update(ok)
# try the same tag會不會被重複insert入DB



# 重設DB的id值:ALTER TABLE User AUTO_INCREMENT = 1; ID 設定是 int，可以存放21億筆資料。
#   設定是bigint unsinged，bigint 的範圍是 -2^63 ~ 2^(63)-1，最大值可以到 18,446,744,073,709,551,615，假設這資料表要用100年，一秒最多可以存50億筆資料
# html標籤清理、格式化輸出、非格式化輸出:https://blog.csdn.net/zhengalen/article/details/51422402
